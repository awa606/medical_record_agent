from __future__ import annotations

import re
import argparse
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = PROJECT_ROOT / "docs" / "final_report" / "AI生成式电子病历辅助系统_期末报告_正式版.md"
REPORT_DOCX = PROJECT_ROOT / "docs" / "final_report" / "AI生成式电子病历辅助系统_期末报告_正式版.docx"


def main() -> int:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError:
        print("python-docx is not installed. Please run: python -m pip install python-docx")
        return 1

    args = parse_args()
    report_md = args.input
    report_docx = args.output

    if not report_md.exists():
        print(f"Report markdown not found: {report_md}")
        return 1

    lines = report_md.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document, Inches, Pt, qn)
    render_cover(document, lines, Pt, qn, WD_ALIGN_PARAGRAPH.CENTER)
    render_markdown(document, body_lines(lines), Pt, qn)

    report_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(report_docx)
    print(report_docx)
    return 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export final report markdown to Word docx.")
    parser.add_argument("--input", type=Path, default=REPORT_MD, help="Markdown report path.")
    parser.add_argument("--output", type=Path, default=REPORT_DOCX, help="Output docx path.")
    return parser.parse_args()


def configure_document(document, inches, pt, qn) -> None:
    section = document.sections[0]
    section.top_margin = inches(1)
    section.bottom_margin = inches(1)
    section.left_margin = inches(1.1)
    section.right_margin = inches(1.1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = pt(12)
    set_style_east_asia(normal, qn, "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = pt(6)

    heading_specs = [
        ("Heading 1", "黑体", 16, 12, 6),
        ("Heading 2", "黑体", 14, 10, 6),
        ("Heading 3", "黑体", 12, 8, 4),
    ]
    for style_name, font_name, size, before, after in heading_specs:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = pt(size)
        style.font.bold = True
        set_style_east_asia(style, qn, font_name)
        style.paragraph_format.space_before = pt(before)
        style.paragraph_format.space_after = pt(after)


def set_style_east_asia(style, qn, font_name: str) -> None:
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style.element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def render_cover(document, lines: list[str], pt, qn, center_alignment) -> None:
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "期末报告")
    rows = cover_rows(lines)

    add_blank(document, 4)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = center_alignment
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, "黑体", 22, qn, bold=True)
    title_paragraph.paragraph_format.space_after = pt(36)

    for label, value in rows:
        paragraph = document.add_paragraph()
        paragraph.alignment = center_alignment
        paragraph.paragraph_format.space_after = pt(10)
        run = paragraph.add_run(f"{label}：{value}")
        set_run_font(run, "宋体", 14, qn)

    add_blank(document, 3)
    note = document.add_paragraph()
    note.alignment = center_alignment
    run = note.add_run("说明：本报告为课程 POC 原型材料，不代表真实临床系统。")
    set_run_font(run, "宋体", 11, qn)
    document.add_page_break()


def add_blank(document, count: int) -> None:
    for _ in range(count):
        document.add_paragraph()


def cover_rows(lines: list[str]) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    in_cover = False
    for line in lines:
        stripped = line.strip()
        if stripped == "## 封面信息":
            in_cover = True
            continue
        if in_cover and stripped.startswith("## "):
            break
        if not in_cover or not stripped.startswith("|"):
            continue
        cells = split_table_row(stripped)
        if len(cells) < 2 or cells[0] == "项目" or is_separator_row(cells):
            continue
        rows.append((cells[0], cells[1]))
    return rows


def body_lines(lines: list[str]) -> list[str]:
    start = 0
    for index, line in enumerate(lines):
        if heading_key(line) == "摘要":
            start = index
            break
    return lines[start:]


def render_markdown(document, lines: list[str], pt, qn) -> None:
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_table(document, table_lines, pt, qn)
            table_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            if in_code:
                add_code_block(document, code_lines, qn)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if is_table_line(stripped):
            table_lines.append(stripped)
            continue

        flush_table()

        if not stripped:
            continue

        if stripped.startswith("## "):
            heading_text = clean_inline(stripped[3:].strip())
            if compact_text(heading_text) in {"摘要", "目录"}:
                add_center_heading(document, heading_text, qn)
            else:
                document.add_heading(heading_text, level=1)
            continue

        if stripped.startswith("### "):
            document.add_heading(clean_inline(stripped[4:].strip()), level=2)
            continue

        if stripped.startswith("#### "):
            document.add_heading(clean_inline(stripped[5:].strip()), level=3)
            continue

        checkbox_match = re.match(r"^- \[( |x|X)\] (.*)$", stripped)
        if checkbox_match:
            add_list_item(document, f"[{checkbox_match.group(1)}] {checkbox_match.group(2)}", qn)
            continue

        if stripped.startswith("- "):
            add_list_item(document, stripped[2:].strip(), qn)
            continue

        ordered_match = re.match(r"^(\d+)\. (.*)$", stripped)
        if ordered_match:
            add_numbered_item(document, ordered_match.group(2).strip(), qn)
            continue

        if stripped.startswith("> "):
            add_body_paragraph(document, stripped[2:].strip(), qn, italic=True, indent=False)
            continue

        is_caption = bool(re.match(r"^图\s*\d+[\s：:]", stripped))
        is_placeholder = stripped.startswith("【截图待插入") or stripped.startswith("【建议插入图")
        add_body_paragraph(
            document,
            stripped,
            qn,
            bold=is_caption,
            italic=is_placeholder,
            indent=not (is_caption or is_placeholder),
        )

    flush_table()
    if in_code and code_lines:
        add_code_block(document, code_lines, qn)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("`", "")


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text)


def heading_key(line: str) -> str:
    stripped = line.strip()
    if not stripped.startswith("#"):
        return ""
    return compact_text(stripped.lstrip("#").strip())


def add_body_paragraph(
    document,
    text: str,
    qn,
    *,
    bold: bool = False,
    italic: bool = False,
    indent: bool = True,
) -> None:
    paragraph = document.add_paragraph()
    if indent:
        from docx.shared import Pt

        paragraph.paragraph_format.first_line_indent = Pt(24)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, "宋体", 12, qn, bold=bold, italic=italic)


def add_center_heading(document, text: str, qn) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 12
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", 16, qn, bold=True)


def add_list_item(document, text: str, qn) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, "宋体", 12, qn)


def add_numbered_item(document, text: str, qn) -> None:
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, "宋体", 12, qn)


def is_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells if cell.strip())


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document, lines: list[str], pt, qn) -> None:
    rows = [split_table_row(line) for line in lines]
    rows = [row for row in rows if not is_separator_row(row)]
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        for column_index in range(column_count):
            value = row_values[column_index] if column_index < len(row_values) else ""
            cell = row.cells[column_index]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(clean_inline(value))
            set_run_font(run, "宋体", 9.5, qn, bold=(row_index == 0))

    document.add_paragraph()


def add_code_block(document, lines: list[str], qn) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 9.5, qn)


def set_run_font(
    run,
    font_name: str,
    size: float,
    qn,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


if __name__ == "__main__":
    raise SystemExit(main())
