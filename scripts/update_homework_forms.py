from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "report"
TEMPLATE_DIR = REPORT_DIR / "appendix_form_templates"
NEW_FROM_V2_DIR = REPORT_DIR / "appendix_form_templates_new_from_v2"
ARCHIVE_DIR = REPORT_DIR / "appendix_form_templates_archive" / "2026-07-08_v1"
HOMEWORK_DIR = PROJECT_ROOT / "homework"
HANDBOOK_V2_SOURCE = Path(
    r"E:\BaiduNetdiskDownload\综合设计3学生成果手册v2.docx"
)
TODAY = date.today().isoformat()
PROJECT_NAME = "Medical Record Agent（AI生成式电子病历辅助系统）"
MEMBER_NAME = "AWA007"


FORM_SPECS = [
    ("01", "个人贡献评价表", "全程表单"),
    ("02", "任务看板模板", "全程表单"),
    ("03", "每日任务记录表_Day1-2", "第1周：需求澄清、系统架构、接口定义"),
    ("04", "功能清单_Day1", "第1周：功能范围与 POC/EVT 划分"),
    ("05", "项目说明_Day1", "第1周：项目概览与评委快速理解材料"),
    ("06", "每日任务记录表_Day3-4", "第1周：核心模块打通"),
    ("07", "系统架构设计表_Day5", "第1周：系统架构设计"),
    ("08", "第1周评分表_POC-Prototype", "第1周：技术打通"),
    ("09", "每日任务记录表_Day6-7", "第2周：系统集成"),
    ("10", "每日任务记录表_Day8-9", "第2周：核心功能实现"),
    ("11", "POC完成检查表_Day10", "第2周：端到端演示"),
    ("12", "第2周评分表_POC-MVP", "第2周：POC达成"),
    ("13", "每日任务记录表_Day11-12", "第3周：性能优化"),
    ("14", "每日任务记录表_Day13-14", "第3周：稳定性测试"),
    ("15", "问题跟踪表_Bug_List_Day15", "第3周：测试复盘"),
    ("16", "系统测试记录表_Day15", "第3周：测试复盘"),
    ("17", "第3周评分表_POC-Refine", "第3周：优化验证"),
    ("18", "每日任务记录表_Day16-17", "第4周：文档完善"),
    ("19", "EVT初级检查表_Day18", "第4周：系统封版"),
    ("20", "第4周评分表_POC交付", "第4周：预答辩"),
    ("21", "综合设计3评分表_总评_Day20", "第4周：终答辩"),
    ("22", "EVT加分表_Day20", "第4周：终答辩"),
]


DAILY_DATA = {
    "Day1-2": {
        "plan": [
            "明确项目从汇报版转向可交付产品，优先打通 MP3/WAV 文件流实时转写。",
            "建立 GitHub Issue、Debug Log、Daily Log、版本记录和开发流程。",
            "梳理 FastAPI、ASR、SSE、医生端前端、知识库和测试入口。",
        ],
        "done": [
            "完成 docs、logs、versions、.github 工程管理结构。",
            "更新 README、architecture、版本演进记录、debug_guide。",
            "创建 v0.2.1 四周迭代计划和 ASR SSE 文件流接口设计。",
        ],
        "blockers": [
            "本机未安装 gh CLI，不能直接命令行创建远端 Issue。",
            "真实医院电脑配置尚未采集。",
        ],
        "support": [
            "已将 Issue Seed 写入 docs/四周迭代执行计划.md，后续可复制到 GitHub Issues。",
            "先以普通医院 Windows PC、CPU-only/集显、16GB 内存作为最低基线。",
        ],
        "next": [
            "实现 ASR session SSE 文件流接口。",
            "在医生端中间转写栏显示分段转写。",
        ],
    },
    "Day3-4": {
        "plan": [
            "新增 ASR session API，支持上传 MP3/WAV 后生成 SSE segment 事件。",
            "复用现有 mock/funasr/qwen3/online ASR 引擎，不改核心算法逻辑。",
            "增加 API 测试，覆盖会话创建、上传、SSE 和结果读取。",
        ],
        "done": [
            "完成 /api/asr/sessions、/audio、/events、/result。",
            "医生端接入 EventSource，实时追加 segment 到中间转写栏。",
            "pytest 全量通过 73 passed，推送 v0.2.1 tag。",
        ],
        "blockers": [
            "当前实时转写为上传后按 ASRResult segment 回放，不是真正底层流式解码。",
        ],
        "support": [
            "文档中明确 v0.2.1 边界，后续真实流式解码放到模型评测和架构升级。",
        ],
        "next": [
            "进入 v0.3 医生/患者角色校正。",
            "补前端逐段角色切换和文本编辑。",
        ],
    },
    "Day6-7": {
        "plan": [
            "实现医生/患者角色校正闭环。",
            "保存校正后重建 conversation_text，并保持旧病历生成接口兼容。",
            "明确当前模型是否需要切换以及后续方言、多语种、多人物路线。",
        ],
        "done": [
            "新增 PATCH /api/asr/sessions/{session_id}/result 保存角色和文本校正。",
            "ASRSegment 新增 needs_review、reviewed_by_doctor、original_text。",
            "医生端中间栏支持角色下拉、文本编辑、保存角色校正。",
            "pytest 全量通过 75 passed，推送 v0.3 tag。",
        ],
        "blockers": [
            "v0.3 仍是结果层人工校正，不是自动 diarization。",
        ],
        "support": [
            "保留人工校正确认作为医生最终边界；多人物自动分离放入 v0.5 评测。",
        ],
        "next": [
            "用真实 FunASR/Qwen3 样本验证角色校正体验。",
            "开始扩展知识库症状-疾病-检查-用药关联规则。",
        ],
    },
    "Day8-9": {
        "plan": [
            "完善主用例：上传音频、实时转写、角色校正、病历生成、医生审核、导出。",
            "优化医生端布局和操作反馈。",
            "补知识库关联规则入口。",
        ],
        "done": [
            "已具备前置能力，待下一轮实现知识库和产品化细节。",
        ],
        "blockers": [
            "知识库仍偏示例化，需要补症状、疾病、检查、用药和规则字段。",
        ],
        "support": [
            "按 docs/四周迭代执行计划.md 创建 feature Issue 后实施。",
        ],
        "next": [
            "进入 v0.4 医学知识库和关联规则扩展。",
        ],
    },
    "Day11-12": {
        "plan": [
            "采集普通医院 PC 配置。",
            "评测 mock、FunASR、Qwen3-ASR、Ollama/本地 LLM。",
            "记录 CER、关键词召回、耗时和资源占用。",
        ],
        "done": [
            "已建立 docs/local_model_edge_benchmark.md 和 docs/asr_model_route.md 评测框架。",
        ],
        "blockers": [
            "真实医院电脑硬件信息和真实音频样本尚未提供。",
        ],
        "support": [
            "先使用普通 Windows PC 基线，后续补独显工作站和边缘端。",
        ],
        "next": [
            "收集硬件配置并运行真实模型评测。",
        ],
    },
    "Day13-14": {
        "plan": [
            "做 30-60 分钟稳定性测试。",
            "验证长音频、错误格式、SSE 断连重连、重复上传和任务失败提示。",
            "完善 Bug List 和 Debug 报告。",
        ],
        "done": [
            "已有单元测试和 API 测试基础，待长时运行验证。",
        ],
        "blockers": [
            "缺少长音频测试样本。",
        ],
        "support": [
            "可先用课程样例音频复制/拼接为长音频进行压力测试。",
        ],
        "next": [
            "执行稳定性测试并记录 Bug List。",
        ],
    },
    "Day16-17": {
        "plan": [
            "完善最终报告、BOM/部署需求、演示脚本和视频材料。",
            "整理 README、版本日志、Debug 记录和测试证据。",
        ],
        "done": [
            "工程管理结构、每日记录和版本目录已建立。",
        ],
        "blockers": [
            "最终演示素材和视频尚未录制。",
        ],
        "support": [
            "在 v1.0 freeze 前集中整理交付包。",
        ],
        "next": [
            "补部署说明和边缘端配置建议。",
        ],
    },
}


def safe_name(name: str) -> str:
    return re.sub(r'[<>:"/\\\\|?*]+', "_", name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    if header and table.rows:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "E8EEF5")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def setup_doc(title: str, subtitle: str | None = None) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.name = "Microsoft YaHei"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    title_p = doc.add_heading(title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    return doc


def add_meta(doc: Document, project: str = PROJECT_NAME, member: str = MEMBER_NAME, day: str = "") -> None:
    rows = [("项目名称", project), ("成员姓名", member), ("日期/阶段", day or TODAY)]
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
    style_table(table)


def add_list_section(doc: Document, heading: str, items: Iterable[str]) -> None:
    doc.add_heading(heading, level=2)
    for item in items:
        doc.add_paragraph(item, style=None)


def add_team_table(doc: Document, contribution: str = "项目开发、文档、测试与版本记录") -> None:
    doc.add_heading("团队成员出勤与贡献", level=2)
    table = doc.add_table(rows=4, cols=3)
    headers = ["成员", "出勤", "今日贡献"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    values = [
        [MEMBER_NAME, "出勤", contribution],
        ["待补充", "待补充", "待补充"],
        ["待补充", "待补充", "待补充"],
    ]
    for r_idx, row in enumerate(values, start=1):
        for c_idx, value in enumerate(row):
            set_cell_text(table.rows[r_idx].cells[c_idx], value)
    style_table(table)
    doc.add_paragraph("项目负责人签字：AWA007")


def add_daily_form(doc: Document, day_key: str, template: bool = False) -> None:
    data = DAILY_DATA.get(day_key, DAILY_DATA["Day1-2"])
    add_meta(doc, day=f"{day_key} / {TODAY}")
    if template:
        data = {
            "plan": ["1.", "2.", "3."],
            "done": ["1.", "2.", "3."],
            "blockers": ["1.", "2.", "3."],
            "support": ["1.", "2."],
            "next": ["1.", "2.", "3."],
        }
    add_list_section(doc, "一、今日计划（早上填写）", data["plan"])
    add_list_section(doc, "二、今日完成（晚上填写）", data["done"])
    add_list_section(doc, "三、问题与阻碍（Blocker）", data["blockers"])
    add_list_section(doc, "四、解决方案 / 需求支持", data["support"])
    add_list_section(doc, "五、明日计划", data["next"])
    add_team_table(doc)


def add_personal_contribution(doc: Document, template: bool = False) -> None:
    add_meta(doc)
    doc.add_heading("一、角色", level=2)
    roles = [
        "☑ 项目负责人(PM)：项目推进、任务分配、进度控制",
        "☑ 系统工程师：系统架构设计、模块协同",
        "☑ 软件/算法工程师：ASR、SSE、角色校正、LLM/病历生成流程",
        "☑ 测试工程师：API 测试、前端语法检查、回归验证",
    ]
    for item in roles if not template else ["□ 项目负责人(PM)", "□ 系统工程师", "□ 软件/算法工程师", "□ 测试工程师"]:
        doc.add_paragraph(item)
    doc.add_heading("二、自评（40%）", level=2)
    doc.add_paragraph("贡献描述：" + ("" if template else "完成工程结构、ASR SSE、角色校正、文档和测试闭环。"))
    doc.add_heading("三、团队互评（30%）", level=2)
    table = doc.add_table(rows=4, cols=3)
    for i, h in enumerate(["评价人", "分数", "评价"]):
        set_cell_text(table.rows[0].cells[i], h, True)
    rows = [["待补充", "待补充", "待补充"], ["待补充", "待补充", "待补充"], ["待补充", "待补充", "待补充"]]
    if not template:
        rows[0] = ["项目自评", "97/100", "已完成原生流式转写、播放器、CAM++ 校准、实时病历预览、病历质量评估、字段证据定位，以及 v0.9.4 候选诊断/治疗建议质量收口与 SSE 断连恢复复验；当前主要待补证据为 HF_TOKEN 放行后的 pyannote measured turns、3D-Speaker 运行区与普通医院 PC 复测。"]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell_text(table.rows[r].cells[c], value)
    style_table(table)
    doc.add_heading("四、导师评价（30%）", level=2)
    doc.add_paragraph("评分：" + ("" if template else "待导师评分"))
    doc.add_paragraph("评价：" + ("" if template else "待导师评价"))
    doc.add_paragraph("最终得分：" + ("" if template else "待汇总"))


def add_kanban(doc: Document, template: bool = False) -> None:
    add_meta(doc, day="任务看板")
    table = doc.add_table(rows=4, cols=3)
    headers = ["To Do（待做）", "Doing（进行中）", "Done（完成）"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    rows = [
        ["pyannote HF_TOKEN 放行与 3D-Speaker 运行区", "v0.9.4 候选诊断与治疗建议质量优化", "v0.9.3 医生端前端冻结与病历质量主线"],
        ["普通医院 PC 复测与 SSE 断连恢复", "pyannote measured turns 与 ASR 对齐复测", "v0.9.2 字段质量、证据定位与质量摘要"],
        ["最终演示视频与导出复盘", "v0.9.5 最终演示数据和稳定性加固", "两说话人 RTTM、v0.8.17-v0.8.20 说话人分离准备、Docker 2626 前端验收"],
    ]
    if template:
        rows = [["", "", ""], ["", "", ""], ["", "", ""]]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell_text(table.rows[r].cells[c], value)
    style_table(table)


def add_feature_list(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="项目组", day="Day1")
    doc.add_paragraph(
        "用途：明确本项目所有功能，划分必须完成（POC）和可选优化（EVT），防止项目范围失控。"
    )
    doc.add_heading("一、功能清单", level=2)
    rows = [["编号", "功能名称", "功能说明", "优先级(P0/P1/P2)", "是否属于POC", "是否属于EVT", "负责人", "完成时间"]]
    rows += [["F01", "", "", "", "", "", "", ""], ["F02", "", "", "", "", "", "", ""], ["F03", "", "", "", "", "", "", ""]] if template else [
        ["F01", "MP3/WAV 上传与 SSE 转写", "上传课程音频后分段显示 ASR 转写结果", "P0", "是", "否", MEMBER_NAME, "v0.2.1"],
        ["F02", "医生/患者角色校正", "逐段切换角色、编辑文本并回写 conversation_text", "P0", "是", "否", MEMBER_NAME, "v0.3"],
        ["F03", "病历生成与医生审核", "生成病历草稿并保留医生审核后导出边界", "P0", "是", "否", MEMBER_NAME, "v0.4"],
        ["F04", "医学知识库关联规则", "输出候选诊断依据、建议检查、风险提醒", "P1", "是", "是", MEMBER_NAME, "v0.4.2"],
        ["F05", "本地多模型 ASR 评测", "对比 FunASR、SenseVoice、Qwen3、Whisper 的 CER/RTF/RSS", "P1", "否", "是", MEMBER_NAME, "v0.5.7"],
        ["F06", "16/30 分钟长音频稳定性", "验证长音频吞吐、内存、CPU、失败处理和部署边界", "P1", "否", "是", MEMBER_NAME, "v0.5.9"],
        ["F07", "医生端产品化与实时病历预览", "优化实时转写区、角色校正区、状态提示，并联动病历草稿与 AI 辅助实时预览", "P1", "否", "是", MEMBER_NAME, "v0.8.5"],
        ["F08", "医生声纹注册与整位角色分类", "登记本机医生声纹，结合 CAM++ 与本地模型做整位说话人角色判定", "P1", "否", "是", MEMBER_NAME, "v0.8.7"],
        ["F09", "说话人分离依赖检查与 RTTM 评测", "检查 pyannote/3D-Speaker/FunASR CAM++ 依赖状态，建立人工 RTTM 评测脚本", "P1", "否", "是", MEMBER_NAME, "v0.8.7"],
        ["F10", "两说话人 RTTM 实测与终答辩冻结", "完成 fever_01、chest_pain_01 人工 RTTM、FunASR CAM++ 评测汇总和 Docker 2601 前端复测", "P1", "否", "是", MEMBER_NAME, "v0.8.8"],
        ["F11", "真实多说话人分离对比与对齐", "补齐 AliMeeting 公开样本、统一比较脚本、ASR-diarization 对齐脚本与真实 failed/skipped 结论", "P1", "否", "是", MEMBER_NAME, "v0.8.18"],
        ["F12", "pyannote 隔离环境与门禁检查", "创建 .venv-diarization、固定 pyannote 依赖版本，并把阻塞收敛到 HF_TOKEN 门禁", "P1", "否", "是", MEMBER_NAME, "v0.8.20"],
        ["F13", "病历质量评估与字段证据定位", "输出字段完整度、低置信度、证据覆盖和建议动作，并把字段证据映射到稳定转写段", "P1", "否", "是", MEMBER_NAME, "v0.9.2"],
        ["F14", "医生端前端冻结与摘要化展示", "默认界面只保留摘要和主操作，完整质量原因、证据和阻断原因进入详情抽屉", "P1", "否", "是", MEMBER_NAME, "v0.9.3"],
    ]
    add_table(doc, rows)
    doc.add_heading("二、优先级建议", level=2)
    add_table(doc, [
        ["P0（必须完成）", "P1（建议完成）", "P2（扩展功能）"],
        [
            "课程结束必须完成，否则 POC 不达标。",
            "完成后项目完整度明显提高，是当前主要迭代对象。",
            "优秀团队或 EVT 方向，视时间和硬件条件推进。",
        ],
    ])


def add_project_brief(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="项目负责人", day="Day1")
    sections = [
        ("一、项目名称", PROJECT_NAME if not template else ""),
        (
            "二、一句话介绍（One Sentence）",
            "" if template else "面向中文医患问诊的 AI 生成式电子病历辅助系统，支持音频转写、角色校正、病历草稿生成和医生审核导出。",
        ),
        (
            "三、解决什么问题（100字以内）",
            "" if template else "将医患问诊录音或文本转为可审核的结构化病历草稿，减少医生重复录入，同时通过候选诊断、知识库规则、日志和审核边界保证工程可追踪。",
        ),
        ("四、目标用户", "" if template else "医生、实习医生、课程评委和医疗 AI 工程演示使用者。"),
        (
            "五、POC目标",
            "" if template else "完成 MP3/WAV 上传、SSE 分段转写、医生/患者角色校正、病历生成、安全审核、导出、医生声纹注册、说话人分离评测入口和本地模型评测证据。",
        ),
        (
            "六、系统组成",
            "" if template else "FastAPI 后端、医生端静态页面、ASR 多引擎适配、MedicalRecordOrchestrator、知识库规则、SQLite/本地日志、测试与报告脚本。",
        ),
        (
            "七、预期成果",
            "" if template else "可运行软件原型、README 与架构文档、每日记录、Debug 记录、模型评测报告、成果手册表单、GitHub 版本记录和最终演示材料。",
        ),
    ]
    for heading, content in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(content or "待填写")


def add_architecture_form(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="项目组", day="Day5")
    add_list_section(doc, "一、系统总体描述", [
        "" if template else "系统以 FastAPI 为后端入口，接收 MP3/WAV 或文本问诊输入，通过 ASR session SSE 推送分段转写，医生在前端校正医生/患者角色后，进入 Agent 编排生成病历草稿，并经医生审核后导出。"
    ])
    doc.add_heading("二、系统架构图（文字版）", level=2)
    doc.add_paragraph("" if template else "static/doctor.html -> /api/asr/sessions -> ASR services -> role correction -> /api/audio/{audio_id}/generate-record -> MedicalRecordOrchestrator -> SQLite/export")
    doc.add_heading("三、模块划分", level=2)
    rows = [["模块名称", "功能描述", "输入", "输出"]]
    rows += [["", "", "", ""], ["", "", "", ""]] if template else [
        ["前端医生工作台", "上传音频、实时查看转写、校正角色、审核病历", "MP3/WAV、文本、医生校正", "ASR 校正结果、审核操作"],
        ["ASR Session API", "管理转写会话、SSE 分段事件、保存角色校正", "音频文件、校正 payload", "ASRResult、conversation_text"],
        ["Agent 编排", "字段抽取、草稿生成、安全校验", "conversation_text", "病历字段、草稿、安全结果"],
    ]
    add_table(doc, rows)
    doc.add_heading("四、接口定义", level=2)
    rows = [["模块A", "模块B", "通信方式", "数据内容"]]
    rows += [["", "", "", ""], ["", "", "", ""]] if template else [
        ["前端", "ASR Session API", "HTTP + SSE", "音频上传、segment events、角色校正"],
        ["ASR Session API", "ASR 引擎", "Python 调用", "audio_id、audio_path、ASRResult"],
        ["音频 API", "Agent 编排", "HTTP/BackgroundTask", "校正后 conversation_text"],
    ]
    add_table(doc, rows)
    add_list_section(doc, "五、关键技术路线", ["1." if template else "1. 文件流优先：先实现上传后 SSE 分段显示。", "2." if template else "2. 医生校正优先：模型角色判断不作为最终事实。", "3." if template else "3. 模型评测后升级：方言、多语种、多人物在 v0.5 决策。"])


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(table.rows[r_idx].cells[c_idx], value, bold=(r_idx == 0))
    style_table(table)


def add_score_form(doc: Document, title: str, total: str, items: list[tuple[str, str]], score: str, template: bool = False) -> None:
    add_meta(doc, day=title)
    for heading, detail in items:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(detail if not template else "□ 高分：\n□ 中等：\n□ 低分：\n□ 0分：")
    doc.add_paragraph(f"总分：{'____' if template else score} / {total}")
    doc.add_paragraph("评委：" + ("" if template else "自评建议，待教师确认"))


def add_poc_check(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="项目组", day="Day10")
    rows = [["功能项", "是否完成", "说明"]]
    rows += [["", "是/否", ""], ["", "是/否", ""]] if template else [
        ["MP3/WAV SSE 实时转写", "是", "已完成 ASR session SSE，支持医生端分段展示。"],
        ["医生/患者角色校正", "是", "已完成逐段校正和保存，可回写 conversation_text。"],
        ["中文医患样本多模型对比", "是", "v0.5.7 已完成 FunASR、SenseVoice、Qwen3 同口径主评测，并补充 CPU/RSS 指标。"],
        ["16/30 分钟长音频稳定性", "是", "v0.5.9 已完成切片稳定性修复与 16/30 分钟样本记录，当前待医院 PC 和真实 ASR 环境复测。"],
        ["医生声纹注册与整位角色分类", "是", "已完成医生声纹注册接口、CAM++ 嵌入匹配和整位说话人角色分类逻辑。"],
        ["两说话人 RTTM 评测与汇总", "是", "已完成 fever_01、chest_pain_01 人工 RTTM、FunASR CAM++ 评测和 summary.md/json。"],
        ["真实多说话人对比与 pyannote 环境", "是", "已完成 v0.8.17-v0.8.20 脚本、报告和 .venv-diarization 安装；当前阻塞为 HF_TOKEN 与 3D-Speaker 运行区。"],
        ["病历质量评估与字段证据定位", "是", "已完成 v0.9.0-v0.9.2 质量评估服务、字段质量状态、segment_id 证据定位与前端质量摘要展示。"],
        ["医生端前端冻结与最终收敛", "是", "已完成 v0.9.3 默认摘要化展示、详情抽屉、导出阻断原因提示和 frontend-freeze 规则。"],
        ["候选诊断/治疗建议质量与 SSE 复验", "是", "已完成 v0.9.4 diagnosis_quality、treatment_safety、建议补问详情展示和 SSE Last-Event-ID 复验；普通医院 PC 复测仍阻塞。"],
    ]
    doc.add_heading("一、核心功能完成情况", level=2)
    add_table(doc, rows)
    add_list_section(doc, "二、系统运行情况", ["☑ 可正常运行" if not template else "□ 可正常运行", "□ 偶发错误", "□ 不稳定"])
    add_list_section(doc, "三、演示情况", ["☑ 可完整演示" if not template else "□ 可完整演示", "□ 部分演示", "□ 无法演示"])
    add_list_section(doc, "四、问题清单（Top5）", [
        "1. 两说话人 RTTM 评测已完成，真实多说话人链路已推进到 pyannote 隔离环境就绪，但 measured turns 仍缺 HF_TOKEN 门禁放行。",
        "2. Qwen3-ASR 可用于角色上下文分类研究，但 CPU-only 资源占用不适合作为默认交付模型。",
        "3. 普通医院 Windows 办公 PC 实机配置与实测数据仍待采集。",
        "4. pyannote 隔离环境已装好但仍需要 HF_TOKEN，3D-Speaker 仍需要独立运行区和 wrapper 脚本。",
        "5. v0.9.3 已冻结前端主界面，v0.9.4 已在冻结约束下补齐候选诊断与治疗建议质量；剩余风险集中在普通医院 PC 复测和多说话人环境阻塞。",
    ] if not template else ["1.", "2.", "3.", "4.", "5."])
    add_list_section(doc, "五、结论", ["☑ 达到阶段性 POC 要求" if not template else "□ 达到POC要求", "□ 未达到POC要求"])


def add_bug_list(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="项目组", day="Day15")
    rows = [["编号", "问题描述", "类型", "严重程度", "状态", "负责人"]]
    rows += [["1", "", "功能/性能/稳定", "高/中/低", "未解决/处理中/已解决", ""], ["2", "", "功能/性能/稳定", "高/中/低", "未解决/处理中/已解决", ""]] if template else [
        ["1", "普通医院 Windows 办公 PC 的真实硬件配置与实测数据尚未采集", "流程", "中", "处理中", MEMBER_NAME],
        ["2", "pyannote 隔离环境已完成安装，但真实 measured 评测仍缺 Hugging Face 令牌授权", "环境", "中", "已记录", MEMBER_NAME],
        ["3", "3D-Speaker 独立运行区、SSE 断连恢复和普通医院 PC 复测仍缺最终验收", "稳定性", "中", "处理中", MEMBER_NAME],
        ["4", "frontend-freeze 生效后，候选诊断和治疗建议质量优化只能通过最小展示适配推进", "流程", "低", "处理中", MEMBER_NAME],
        ["5", "FunASR 首次冷启动和模型下载可能导致前端等待首段超过 150 秒", "环境", "中", "已记录", MEMBER_NAME],
    ]
    add_table(doc, rows)
    add_list_section(doc, "说明", ["高 = 系统不能运行/影响系统运行", "中 = 功能受影响/影响体验", "低 = 优化项"])


def add_test_record(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="测试负责人", day="Day15")
    add_list_section(doc, "一、测试目标", ["验证成果手册表单刷新、v0.9.4 候选诊断/治疗建议质量主线未回退、v0.9.3 前端冻结状态可追踪，并补 SSE 断连恢复复验证据。"] if not template else ["（如稳定性/精度/功耗）"])
    doc.add_heading("二、测试方案", level=2)
    rows = [["测试项", "方法", "指标"]]
    rows += [["", "", ""], ["", "", ""]] if template else [
        ["说话人链路定向回归", "pytest -q tests/test_diarization_evaluator.py tests/test_speaker_profiles.py tests/test_speaker_role_classifier.py tests/test_asr_sessions_api.py tests/test_funasr_streaming_engine.py tests/test_records_api.py tests/test_speaker_diarization.py tests/test_run_diarization_engine_compare.py tests/test_apply_diarization_turns_to_asr_result.py tests/test_bootstrap_diarization_eval.py", "37 passed"],
        ["病历质量定向回归", "pytest -q tests/test_record_quality.py tests/test_records_api.py", "通过"],
        ["依赖与语法检查", "py_compile + node --check", "通过"],
        ["说话人分离依赖报告", "python scripts/check_diarization_dependencies.py && .\\.venv-diarization\\Scripts\\python scripts\\bootstrap_diarization_eval.py", "生成 dependency_status.md/json 和 bootstrap_status.md/json"],
    ]
    add_table(doc, rows)
    doc.add_heading("三、测试结果", level=2)
    rows = [["测试项", "结果", "是否达标"]]
    rows += [["", "", "是/否"], ["", "", "是/否"]] if template else [
        ["表单脚本与定向回归", "py_compile + pytest 定向用例 + update_homework_forms.py", "是"],
        ["说话人链路定向测试", "37 passed（含 v0.8.17-v0.8.20 新增脚本用例）", "是"],
        ["病历质量定向测试", "tests/test_record_quality.py 与 tests/test_records_api.py 通过", "是"],
        ["两说话人 RTTM 评测", "fever_01/chest_pain_01 measured，summary.md/json 已生成", "是"],
        ["Docker 2626 前端复测", "文本、Mock ASR、质量摘要、详情抽屉与导出阻断原因已验收", "是"],
    ]
    add_table(doc, rows)
    add_list_section(doc, "四、问题分析", ["v0.9.4 候选诊断/治疗建议质量和 SSE 断连恢复复验证据已具备；普通医院 PC 实机复测、HF_TOKEN 门禁放行和 3D-Speaker 运行区仍未完成，且前端冻结后只能做最小展示改动。"] if not template else ["1.", "2."])
    add_list_section(doc, "五、改进措施", ["按 v0.9.5 计划在 frontend-freeze 约束下优先补普通医院 PC 复测、pyannote measured turns、ASR 对齐复测和最终演示视频；提交时排除模型缓存、运行时音频、HF_TOKEN 和私密配置。"] if not template else ["1.", "2."])


def add_evt_check(doc: Document, template: bool = False) -> None:
    add_meta(doc, member="项目组", day="Day18")
    items = [
        ("一、样机数量", ["☑ 1套软件原型" if not template else "□ 1台", "□ 2台", "□ ≥3台"]),
        ("二、一致性情况", ["☑ 基本一致" if not template else "□ 完全一致", "□ 差异较大"]),
        ("三、测试情况", ["☑ 有部分测试" if not template else "□ 有完整测试", "□ 无测试"]),
        ("四、FMEA（简化）", ["□ 有", "☑ 待补充" if not template else "□ 无"]),
        ("五、结论", ["☑ 达到软件 POC 阶段要求，EVT 硬件加分待后续补充" if not template else "□ 达到EVT初级"]),
    ]
    for heading, values in items:
        add_list_section(doc, heading, values)


def add_final_score(doc: Document, template: bool = False) -> None:
    add_meta(doc, day="Day20")
    rows = [["类别", "分值", "当前填写"]]
    rows += [["POC实现", "50", ""], ["工程能力", "20", ""], ["文档", "15", ""], ["展示答辩", "15", ""]] if template else [
        ["POC实现", "50", "已完成 SSE、角色校正、病历生成闭环、病历质量评估、字段证据定位，以及 v0.9.4 候选诊断/治疗建议质量收口与 SSE 复验；保留 v0.8.8 与 v0.8.17-v0.8.20 说话人分离评测证据，建议 49/50。"],
        ["工程能力", "20", "工程结构、测试、版本、日志、评测脚本和表单自动化完整，且已形成 frontend-freeze 规则与质量回归闭环，建议 19/20。"],
        ["文档", "15", "README、版本、架构、模型路线、日报、能力矩阵和周评审材料已更新，建议 14/15。"],
        ["展示答辩", "15", "前端冻结说明、验收文档、runbook 和周评审材料已补齐，待录制最终演示视频并补普通医院 PC 复测，暂填 13/15。"],
    ]
    add_table(doc, rows)
    doc.add_paragraph("总分：" + ("" if template else "建议自评 95/100，待教师确认"))
    doc.add_paragraph("评委：" + ("" if template else "待教师评分"))


def add_evt_bonus(doc: Document, template: bool = False) -> None:
    add_meta(doc, day="Day20")
    rows = [["项目", "最高分", "当前情况"]]
    rows += [["多样机", "5", ""], ["测试能力", "5", ""], ["FMEA", "5", ""], ["工程优化", "5", ""]] if template else [
        ["多样机", "5", "当前为软件原型，暂不申请多样机加分。"],
        ["测试能力", "5", "已有 pytest、JS 检查、服务烟测、Qwen3 同口径补测、16/30 分钟长音频稳定性测试入口、两说话人 RTTM 实测、真实多说话人分离对比/对齐脚本和 Docker 2601 前端验收。"],
        ["FMEA", "5", "已在冻结清单中列出真实患者数据、API Key、模型缓存、大文件、HF_TOKEN/3D-Speaker 门禁和普通医院 PC 待补等风险边界。"],
        ["工程优化", "5", "已完成 SSE、角色校正、证据回放、评测脚本、pyannote 隔离环境准备、表单自动化和冻结清单，建议申请部分加分。"],
    ]
    add_table(doc, rows)
    doc.add_paragraph("加分：" + ("" if template else "建议 3-5/20，待教师确认"))


def build_form_doc(code: str, name: str, stage: str, template: bool = False) -> Document:
    doc = setup_doc(name, f"{stage} | {'空白模板' if template else '已填写'} | {TODAY}")
    if "个人贡献" in name:
        add_personal_contribution(doc, template)
    elif "任务看板" in name:
        add_kanban(doc, template)
    elif "功能清单" in name:
        add_feature_list(doc, template)
    elif "项目说明" in name:
        add_project_brief(doc, template)
    elif "每日任务记录表" in name:
        match = re.search(r"Day(\\d+)-(\\d+)", name)
        day_key = f"Day{match.group(1)}-{match.group(2)}" if match else "Day1-2"
        add_daily_form(doc, day_key, template)
    elif "系统架构设计" in name:
        add_architecture_form(doc, template)
    elif "第1周评分" in name:
        add_score_form(doc, "第1周评分表：POC-Prototype", "20", [
            ("一、技术打通（8分）", "☑ 7-8分：ASR session、SSE、前端显示、工程日志等核心模块已真实运行并联通。"),
            ("二、系统架构设计（6分）", "☑ 5-6分：架构文档、接口和模块划分已完成。"),
            ("三、接口定义（3分）", "☑ 3分：ASR session、audio、tasks API 已明确。"),
            ("四、过程记录（3分）", "☑ 3分：logs/daily、版本演进记录、能力证据追踪矩阵 已维护。"),
        ], "20", template)
    elif "POC完成检查" in name:
        add_poc_check(doc, template)
    elif "第2周评分" in name:
        add_score_form(doc, "第2周评分表：POC-MVP", "30", [
            ("一、核心功能实现（15分）", "☑ 核心链路可运行：上传、转写、角色校正、病历生成。"),
            ("二、系统集成（8分）", "☑ 前后端、ASR、Agent、日志已集成。"),
            ("三、实体原型（5分）", "☑ 软件原型完整，硬件实体不适用。"),
            ("四、基础测试（2分）", "☑ pytest 和前端语法检查通过。"),
        ], "28", template)
    elif "Bug" in name or "问题跟踪" in name:
        add_bug_list(doc, template)
    elif "系统测试记录" in name:
        add_test_record(doc, template)
    elif "第3周评分" in name:
        add_score_form(doc, "第3周评分表：POC-Refine", "20", [
            ("一、系统稳定性（8分）", "已完成长音频切片稳定性记录、两说话人 RTTM 评测和 Docker 2601 前端复测；普通医院 PC 待补。"),
            ("二、性能优化（6分）", "已补 FunASR/SenseVoice/Qwen3 对比和前端冷启动边界记录。"),
            ("三、测试完整性（4分）", "已有 API/单元测试、JS 检查、RTTM 评测脚本和前端截图验收。"),
            ("四、问题闭环（2分）", "已建立 Bug List、Debug 记录和 v1.0 冻结清单。"),
        ], "待评", template)
    elif "EVT初级" in name:
        add_evt_check(doc, template)
    elif "第4周评分" in name:
        add_score_form(doc, "第4周评分表：POC交付", "30", [
            ("一、POC完整性（10分）", "原型、文档、版本记录、表单和 v1.0 冻结清单已具备。"),
            ("二、工程质量（10分）", "代码结构、测试、日志、评测脚本和证据矩阵已建立。"),
            ("三、展示与答辩（10分）", "PPT 大纲、讲稿、runbook 和截图证据已补齐，最终视频待录制。"),
        ], "27", template)
    elif "综合设计3评分" in name:
        add_final_score(doc, template)
    elif "EVT加分" in name:
        add_evt_bonus(doc, template)
    else:
        add_meta(doc)
        doc.add_paragraph("待填写")
    return doc


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_index_doc(template: bool = True) -> Document:
    title = "学术成果手册附录表单模板汇总" if template else "Medical Record Agent 学术成果表单填写汇总"
    doc = setup_doc(title, f"生成日期：{TODAY}")
    doc.add_paragraph("来源：report/综合设计3学生成果手册v2.docx")
    for code, name, stage in FORM_SPECS:
        doc.add_page_break()
        doc.add_heading(f"{code}. {name}", level=1)
        doc.add_paragraph(stage)
        form_doc = build_form_doc(code, name, stage, template=template)
        for element in form_doc.element.body:
            if element.tag.endswith("sectPr"):
                continue
            doc.element.body.append(element)
    return doc


def archive_existing_templates() -> None:
    if ARCHIVE_DIR.exists() or not TEMPLATE_DIR.exists():
        return
    existing_docs = list(TEMPLATE_DIR.glob("*.docx"))
    if not existing_docs:
        return
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    for path in existing_docs:
        shutil.copy2(path, ARCHIVE_DIR / path.name)


def clean_generated_docs(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    for old in folder.glob("*.docx"):
        old.unlink()


def resolve_handbook_source() -> Path | None:
    for candidate in [
        HANDBOOK_V2_SOURCE,
        REPORT_DIR / "综合设计3学生成果手册v2.docx",
        REPORT_DIR / "综合设计3学生成果手册.docx",
    ]:
        if candidate.exists():
            return candidate
    return None


def main() -> None:
    archive_existing_templates()
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    NEW_FROM_V2_DIR.mkdir(parents=True, exist_ok=True)
    HOMEWORK_DIR.mkdir(parents=True, exist_ok=True)
    (HOMEWORK_DIR / "individual_forms").mkdir(parents=True, exist_ok=True)

    # Clean generated DOCX files only.
    for folder in [TEMPLATE_DIR, NEW_FROM_V2_DIR, HOMEWORK_DIR / "individual_forms"]:
        clean_generated_docs(folder)

    save_doc(build_index_doc(template=True), TEMPLATE_DIR / "00_附录表单模板汇总.docx")
    save_doc(build_index_doc(template=False), HOMEWORK_DIR / "00_Medical_Record_Agent_学术成果表单填写汇总.docx")

    for code, name, stage in FORM_SPECS:
        save_doc(
            build_form_doc(code, name, stage, template=True),
            TEMPLATE_DIR / f"{code}_{safe_name(name)}_模板.docx",
        )
        save_doc(
            build_form_doc(code, name, stage, template=False),
            HOMEWORK_DIR / "individual_forms" / f"{code}_{safe_name(name)}_已填写.docx",
        )
        if name in {"功能清单_Day1", "项目说明_Day1"}:
            save_doc(
                build_form_doc(code, name, stage, template=True),
                NEW_FROM_V2_DIR / f"{code}_{safe_name(name)}_新增表单模板.docx",
            )

    source = resolve_handbook_source()
    if source is not None:
        shutil.copy2(source, REPORT_DIR / "学术成果手册v2_源文件备份.docx")


if __name__ == "__main__":
    main()
